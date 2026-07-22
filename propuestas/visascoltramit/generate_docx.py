#!/usr/bin/env python3
"""Generate editable DOCX proposal matching Paola Hoyos visual identity."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

# Brand colors (Paola)
PINK = RGBColor(0x94, 0x38, 0x50)
PINK_DEEP = RGBColor(0x7A, 0x2F, 0x42)
INK = RGBColor(0x2D, 0x1E, 0x1B)
BODY = RGBColor(0x4A, 0x3E, 0x3C)
MUTED = RGBColor(0x55, 0x49, 0x47)
FOOTER = RGBColor(0x23, 0x18, 0x16)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SOFT = RGBColor(0xE8, 0xA7, 0xB3)
LINE = "E8E4DE"
PINK_BG = "FFF0F2"
WARM_BG = "F3E6E0"
CREAM = "FAF7F5"
DARK = "231816"


def set_run_font(run, name="Calibri", size=10, bold=False, italic=False, color=BODY):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shd)


def set_cell_border(cell, color="E8E4DE"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:color="{color}"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(borders)


def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.clear() if False else None


def add_para(doc_or_cell, text, *, size=10, bold=False, italic=False, color=BODY,
             font="Calibri", align="left", space_before=0, space_after=8, serif=False):
    if hasattr(doc_or_cell, "paragraphs") and hasattr(doc_or_cell, "add_paragraph"):
        # Document or cell
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.35
    run = p.add_run(text)
    fname = "Georgia" if serif else font
    set_run_font(run, name=fname, size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich(doc, parts, *, align="left", space_before=0, space_after=8, serif=False):
    """parts: list of (text, kwargs)"""
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    for text, kw in parts:
        run = p.add_run(text)
        fname = "Georgia" if (serif or kw.get("serif")) else kw.get("font", "Calibri")
        set_run_font(
            run,
            name=fname,
            size=kw.get("size", 10),
            bold=kw.get("bold", False),
            italic=kw.get("italic", False),
            color=kw.get("color", BODY),
        )
    return p


def eyebrow(doc, text):
    return add_para(doc, text.upper(), size=8, bold=True, color=PINK, space_after=6)


def h1(doc, text):
    return add_para(doc, text, size=22, color=INK, serif=True, space_before=2, space_after=8)


def h2(doc, text):
    return add_para(doc, text, size=14, bold=True, color=INK, space_before=12, space_after=6)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10, bold=True, color=INK)
        r2 = p.add_run(text)
        set_run_font(r2, size=10, color=BODY)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10, color=BODY)
    return p


def quote(doc, text):
    p = add_para(doc, text, size=12, italic=True, color=INK, serif=True, space_before=10, space_after=12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="10" w:color="943850"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)
    return p


def add_table_card(doc, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows_data):
        for j, cell_info in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            text = cell_info.get("text", "")
            bg = cell_info.get("bg")
            color = cell_info.get("color", BODY)
            bold = cell_info.get("bold", False)
            size = cell_info.get("size", 9)
            if bg:
                shade_cell(cell, bg)
            set_cell_border(cell, "E8E4DE")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            set_run_font(run, size=size, bold=bold, color=color if not isinstance(color, str) else RGBColor.from_string(color) if False else color)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # ---- COVER ----
    # Dark cover simulation via shaded full-width table
    cover = doc.add_table(rows=1, cols=1)
    cell = cover.cell(0, 0)
    shade_cell(cell, DARK)
    cell.width = Cm(17.4)

    def cover_p(text, **kw):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(kw.get("sb", 0))
        p.paragraph_format.space_after = Pt(kw.get("sa", 6))
        if kw.get("align") == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        set_run_font(
            run,
            name="Georgia" if kw.get("serif") else "Calibri",
            size=kw.get("size", 10),
            bold=kw.get("bold", False),
            italic=kw.get("italic", False),
            color=kw.get("color", WHITE),
        )
        return p

    # clear default empty para
    cell.paragraphs[0].clear()
    cover_p("PAOLA HOYOS  ·  CREATIVE", size=9, bold=True, color=SOFT, sa=2)
    cover_p("Propuesta comercial  ·  Julio 2026  ·  Confidencial", size=8, color=RGBColor(0xD8, 0xC4, 0xC9), sa=20)
    cover_p("PRODUCCIÓN & EDICIÓN DE CONTENIDO", size=8, bold=True, color=SOFT, sa=10)
    cover_p("Transformamos sus guiones en contenido que inspira confianza.", size=22, serif=True, color=WHITE, sa=12)
    cover_p(
        "Propuesta de producción audiovisual para Instagram: dirección de grabación, "
        "edición creativa y piezas listas para publicar — con la autoridad visual que VisasColTramit merece.",
        size=10,
        color=RGBColor(0xD8, 0xC4, 0xC9),
        sa=28,
    )
    cover_p("Preparado para", size=8, color=RGBColor(0xD8, 0xC4, 0xC9), sa=0)
    cover_p("VisasColTramit", size=12, bold=True, color=WHITE, sa=10)
    cover_p("Preparado por", size=8, color=RGBColor(0xD8, 0xC4, 0xC9), sa=0)
    cover_p("Paola Hoyos Cardona", size=12, bold=True, color=WHITE, sa=8)

    page_break(doc)

    # ---- 01 ENTENDIMIENTO ----
    eyebrow(doc, "01 · Entendimiento de marca")
    h1(doc, "VisasColTramit, en su esencia")
    add_para(
        doc,
        "No se trata solo de trámites. Se trata de acompañar proyectos de vida — familias, "
        "profesionales y nómadas digitales — con rigor legal y calidez humana.",
        size=11,
        color=MUTED,
        space_after=12,
    )
    h2(doc, "Propuesta de valor")
    add_para(
        doc,
        "Orientación clara, servicio bilingüe y acompañamiento de principio a fin en visas "
        "colombianas, pasaportes de EE.UU. y Canadá. El diferencial no es solo el documento: "
        "es la certeza de no caminar solo.",
    )
    h2(doc, "Público objetivo")
    for t in [
        "Extranjeros que quieren vivir o invertir en Colombia",
        "Nómadas digitales y estudiantes internacionales",
        "Familias en procesos de agrupación / residencia",
        "Profesionales con pasaporte EE.UU. o Canadá a renovar",
    ]:
        bullet(doc, t)
    h2(doc, "ADN de comunicación")
    bullet(doc, " especialistas, procesos estructurados", bold_prefix="Autoridad:")
    bullet(doc, " cumplimiento, transparencia, rigor", bold_prefix="Confianza:")
    bullet(doc, " empatía, acompañamiento humano", bold_prefix="Cercanía:")
    bullet(doc, " lenguaje simple frente a la burocracia", bold_prefix="Claridad:")
    h2(doc, "Identidad visual observada")
    add_para(
        doc,
        "Verde esmeralda institucional, blanco limpio, acentos dorados/amarillos (referencia Colombia), "
        "tipografía sans-serif bold en overlays. Contenido mixto: talking-heads + gráficos informativos.",
    )
    quote(
        doc,
        "“Detrás de cada trámite hay un proyecto de vida.” Esa frase es el norte creativo: "
        "el contenido debe verse tan sólido como el proceso que ofrecen.",
    )

    page_break(doc)

    # ---- 02 OPORTUNIDAD ----
    eyebrow(doc, "02 · Lectura creativa")
    h1(doc, "La oportunidad no es inventar temas. Es elevar la ejecución.")
    add_para(
        doc,
        "Ustedes ya tienen los guiones y la planeación. El salto que falta es de producción: "
        "dirección, ritmo, consistencia visual y acabado editorial que proyecte la misma seriedad "
        "de una firma migratoria de referencia.",
        size=11,
        color=MUTED,
    )
    h2(doc, "Lo que ya funciona")
    for t in [
        "Temática educativa clara (visas, checklists, pasos)",
        "Presencia humana en cámara (rostro = confianza)",
        "Uso del verde de marca como ancla visual",
        "Tono útil, no intimidante",
    ]:
        bullet(doc, t)
    h2(doc, "Lo que vamos a elevar")
    for t in [
        "Iluminación, encuadre y continuidad entre piezas",
        "Ritmo de edición (hooks, cuts, retención)",
        "Motion graphics y subtítulos con sistema tipográfico",
        "B-roll que humanice sin perder profesionalismo",
    ]:
        bullet(doc, t)
    h2(doc, "Promesa de esta propuesta")
    add_para(
        doc,
        "Convertir cada guion en una pieza profesional que se vea — y se sienta — como una marca "
        "que ya es referente: segura, cercana y visualmente impecable. Sin rehacer su estrategia. "
        "Potenciando su ejecución.",
        bold=False,
    )

    page_break(doc)

    # ---- 03 SERVICIO ----
    eyebrow(doc, "03 · El servicio")
    h1(doc, "Producción y edición de contenido para Instagram")
    add_para(
        doc,
        "Un sistema mensual de piezas listas para publicar. Ustedes aportan los guiones y la agenda "
        "temática; yo dirijo, produzco y edito con estándar premium.",
        size=11,
        color=MUTED,
    )
    add_para(doc, "INVERSIÓN MENSUAL", size=8, bold=True, color=PINK, space_before=8, space_after=2)
    add_para(doc, "$600.000 COP / mes", size=26, color=INK, serif=True, space_after=4)
    add_para(
        doc,
        "12 publicaciones · Reels (máx. 40 s) y/o carruseles · Dirección + producción + edición creativa",
        size=9,
        color=MUTED,
    )
    h2(doc, "Qué incluye, en concreto")
    items = [
        ("Preproducción", "Brief de rodaje por pieza a partir de sus guiones: plano, hook visual, B-roll, vestuario y referencias de ritmo."),
        ("Dirección en set", "Acompañamiento creativo durante la grabación (presencial o remoto guiado) para asegurar autoridad, cercanía y calidad de imagen."),
        ("Producción AV", "Selección de tomas, organización de rushes, criterios de continuidad y captura orientada a edición vertical 9:16."),
        ("Postproducción", "Edición, subtítulos quemados, tipografías de marca, lower-thirds, stickers tipográficos, música libre de derechos, colorización y exportación lista para Instagram."),
        ("Entrega", "Archivos finales en formato vertical + versión carrusel cuando aplique. Una ronda de ajustes por pieza."),
    ]
    for title, desc in items:
        add_para(doc, title, size=11, bold=True, color=INK, space_before=6, space_after=2)
        add_para(doc, desc, size=10, color=BODY, space_after=4)

    page_break(doc)

    # ---- 04 DIRECCIÓN ----
    eyebrow(doc, "04 · Dirección de grabación")
    h1(doc, "En set, cada detalle construye autoridad")
    add_para(
        doc,
        "La diferencia entre un video “correcto” y uno que convierte está en la dirección: "
        "cómo se mira a cámara, cómo se ocupa el cuadro, cómo se respira el mensaje.",
        size=11,
        color=MUTED,
    )
    h2(doc, "Presencia y lenguaje corporal")
    for t in [
        "Postura abierta, hombros relajados, pecho ligeramente hacia lente",
        "Gesto con las manos dentro del cuadro (nunca cortar muñecas)",
        "Mirada directa al lente en hooks; mirada “off” en storytelling",
        "Pausas conscientes de 0.3–0.5 s antes de la frase clave",
        "Sonrisa contenido / seriedad cálida en temas legales",
    ]:
        bullet(doc, t)
    h2(doc, "Vestuario y styling")
    for t in [
        "Bloque sólido: negro, blanco, beige o verde marca",
        "Evitar rayas finas y logos ajenos (moiré / ruido visual)",
        "Accesorio mínimo: reloj fino o aretes discretos",
        "Look consistente por bloque de 4–6 piezas del mes",
    ]:
        bullet(doc, t)
    h2(doc, "Iluminación")
    for t in [
        "Key light a 45° (ventana suave o softbox)",
        "Fill ligero para no hundir ojos",
        "Evitar luz cenital dura de oficina",
        "Fondo 1–1.5 stops más oscuro que el rostro",
        "Balance de blancos fijo (no auto) por sesión",
    ]:
        bullet(doc, t)
    h2(doc, "Audio y entorno")
    for t in [
        "Micrófono solapa o close-mic al celular",
        "Ambiente controlado: sin eco de salas vacías",
        "Fondo limpio: planta, estantería ordenada o muro neutro",
        "Profundidad: sin papeles sueltos ni pantallas brillantes",
    ]:
        bullet(doc, t)
    quote(
        doc,
        "Dirigir no es “decir qué decir”. Es diseñar cómo se siente la marca cuando alguien "
        "escucha un consejo migratorio por primera vez.",
    )

    page_break(doc)

    # ---- 05 CÓMO GRABAR ----
    eyebrow(doc, "05 · Cómo grabar cada tipo de video")
    h1(doc, "Kit de encuadres para Instagram 9:16")
    add_para(
        doc,
        "Un sistema práctico de planos, ángulos y movimientos para que cada guion se grabe "
        "con intención — no improvisando el cuadro.",
        size=11,
        color=MUTED,
    )

    shots = [
        ("Talking head experto", "Medio corto (pecho-cabeza) + detalle de manos. Frontal a altura de ojos. Push-in lento al CTA. Rostro en tercio superior; espacio a la derecha para texto."),
        ("Hook de duda / mito", "Primer plano expresivo → medio. Ligeramente contrapicado. Whip-pan a B-roll. Texto grande arriba; sujeto centrado 1.5 s."),
        ("Checklist / pasos", "Medio + inserts de documentos/pantalla/mapa. Estático limpio; cortes cada 2–3 s. Lower-third numerado."),
        ("Testimonio / caso", "Medio lateral ¾ + B-roll emocional. Ángulo 15–30°. Travelling suave. Luz cálida; menos gráficos."),
        ("Carrusel educativo", "Portada foto + tipografía; interiores limpios. Foto hero con profundidad. Margen seguro; tipografía en 3 niveles."),
    ]
    for title, desc in shots:
        add_para(doc, title, size=11, bold=True, color=PINK, space_before=8, space_after=2)
        add_para(doc, desc, size=10, space_after=4)

    h2(doc, "Reglas de oro en cámara")
    for t in [
        "Vertical nativo 1080×1920 (nunca crop de horizontal)",
        "Lente equivalente 24–35 mm (celular wide sin ultra-wide)",
        "Estabilización: trípode o gimbal; nada de temblor “home video”",
        "Grabar 3 tomas del hook + 2 del cuerpo + 1 CTA limpio",
        "Dejar 1 s de cola al inicio/fin para edición",
        "Claqueta verbal: nombre de pieza + take",
    ]:
        bullet(doc, t)

    page_break(doc)

    # ---- 06 EDICIÓN ----
    eyebrow(doc, "06 · Edición para el algoritmo")
    h1(doc, "Ritmo, retención y acabado editorial")
    add_para(
        doc,
        "Instagram premia retención y rewatch. La edición no es “cortar bonito”: es diseñar "
        "la atención segundo a segundo, sin traicionar el tono serio de la marca.",
        size=11,
        color=MUTED,
    )
    h2(doc, "Arquitectura de un Reel ≤ 40 s")
    bullet(doc, " hook visual + texto que detiene el scroll", bold_prefix="0–1.5 s:")
    bullet(doc, " promesa o tensión (mito, error, plazo)", bold_prefix="1.5–8 s:")
    bullet(doc, " valor en bloques de 3–5 s con cuts", bold_prefix="8–30 s:")
    bullet(doc, " síntesis + CTA claro (WhatsApp / link)", bold_prefix="30–40 s:")
    h2(doc, "Subtítulos y tipografía")
    for t in [
        "Subtítulos quemados, palabra a palabra o por frase corta",
        "Máx. 2 líneas; contraste alto sobre verde/blanco/negro",
        "Highlight de keywords en verde marca o dorado",
        "Safe zone: lejos de UI de Instagram (abajo/derecha)",
    ]:
        bullet(doc, t)
    h2(doc, "Motion graphics sencillos")
    for t in [
        "Lower-thirds con nombre / rol",
        "Numeración animada de pasos (01, 02, 03)",
        "Checkmarks y sellos “aprobado / importante”",
        "Transiciones: cut seco, match-cut, zoom sutil 3–5%",
        "Sin efectos baratos ni stickers genéricos",
    ]:
        bullet(doc, t)
    h2(doc, "Música, color y sonido")
    for t in [
        "Tracks trending o corporativos suaves (bajo volumen bajo voz)",
        "Duck music −12 a −18 dB bajo locución",
        "Color: piel natural, verdes un punto más ricos, negros limpios",
        "SFX discretos en cortes de lista (whoosh / tick)",
    ]:
        bullet(doc, t)
    quote(
        doc,
        "Principio de edición VisasColTramit: dinámico, sí. Ansioso, no. El ritmo debe transmitir "
        "control — como un proceso migratorio bien llevado: claro, ordenado y humano.",
    )

    page_break(doc)

    # ---- 07 IDENTIDAD ----
    eyebrow(doc, "07 · Identidad visual en contenido")
    h1(doc, "Consistencia que se reconoce antes de leer el logo")
    add_para(
        doc,
        "Cada Reel y carrusel debe pertenecer a la misma familia visual. Eso es lo que convierte "
        "una cuenta nueva en una marca memorable.",
        size=11,
        color=MUTED,
    )
    h2(doc, "Sistema gráfico")
    for t in [
        "Verde institucional como color de acento y highlights",
        "Dorado/ámbar solo para “alerta / plazo / bandera”",
        "Fondos: blanco editorial o verde profundo en portadas",
        "Esquinas redondeadas suaves en cards internas",
        "Iconografía lineal simple (pasaporte, check, mapa)",
    ]:
        bullet(doc, t)
    h2(doc, "Tipografía en pantalla")
    for t in [
        "Hook: bold condensed / display corto",
        "Cuerpo: sans geométrica limpia",
        "CTA: botón tipográfico o pill verde",
        "Máximo 3 pesos por pieza",
    ]:
        bullet(doc, t)
    h2(doc, "Percepción deseada: autoridad + confianza + cercanía")
    add_para(doc, "Autoridad — encuadres estables, data clara, voz firme, tipografía sobria.", space_after=4)
    add_para(doc, "Confianza — continuidad visual, subtítulos impecables, cero improvisación estética.", space_after=4)
    add_para(doc, "Cercanía — sonrisa genuina, B-roll humano, lenguaje accesible, CTA cálido (“te acompañamos”).", space_after=8)
    h2(doc, "Grid Instagram")
    for t in [
        "Alternancia consciente: Reel humano → carrusel → Reel checklist",
        "Portadas de Reel con tipografía consistente",
        "Misma familia de thumbnails = feed premium",
    ]:
        bullet(doc, t)

    page_break(doc)

    # ---- 08 EJEMPLOS ----
    eyebrow(doc, "08 · Ejemplos de producción")
    h1(doc, "De guion a pieza: casos de dirección")
    add_para(
        doc,
        "Ejemplos aplicados al universo VisasColTramit. Sirven como referencia de cómo se traduce "
        "un objetivo en planos, B-roll, ritmo gráfico y CTA.",
        size=11,
        color=MUTED,
    )

    examples = [
        {
            "title": "Ejemplo 01 · Reel 25–35 s — “Antes de aplicar a tu visa de estudiante”",
            "items": [
                ("Objetivo", "Educar y captar leads calificados con dudas reales de visa V."),
                ("Tipo de plano", "Hook en primer plano → medio corto → detalle de documentos."),
                ("Ángulos", "Frontal a ojos + inserto top-down de checklist impresa."),
                ("B-roll", "Pasaporte abriéndose, laptop con formulario, skyline Medellín, firma de papeles."),
                ("Ritmo", "Corte cada 2.5–3.5 s. Zoom 4% en cada tip. Música corporate-upbeat suave."),
                ("Gráficos", "Números 01–03 en verde; highlight de “error común”; lower-third “Visa estudiantes”."),
                ("CTA", "“Agenda tu asesoría de 10 min — link en bio / WhatsApp”."),
            ],
        },
        {
            "title": "Ejemplo 02 · Reel 20–30 s — “Checklist nómada digital”",
            "items": [
                ("Objetivo", "Alcance + guardados. Pieza ultra útil y rewatchable."),
                ("Tipo de plano", "Talking head medio + jump cuts a inserts de ítems."),
                ("Ángulos", "Frontal limpio; cámara fija; B-roll lateral café/cowork."),
                ("B-roll", "Manos en teclado, café, visa stamp, maleta, mapa de Colombia."),
                ("Ritmo", "Staccato: 1 ítem cada 2 s. SFX tick al check."),
                ("Gráficos", "Checklist animada, checks verdes, tipografía bold en hooks."),
                ("CTA", "“Guarda este video y escríbenos ‘NÓMADA’ por WhatsApp”."),
            ],
        },
        {
            "title": "Ejemplo 03 · Reel 30–40 s — “Traspaso de visa residente: hazlo antes de octubre”",
            "items": [
                ("Objetivo", "Urgencia responsable + conversión a trámite de traspaso."),
                ("Tipo de plano", "Medio corto serio; cutaways a calendario y documento Tipo R."),
                ("Ángulos", "Ligeramente contrapicado (autoridad). Hold largo en la fecha límite."),
                ("B-roll", "Close-up de visa, hoja oficial, calendario, firma, handshake."),
                ("Ritmo", "Hook tenso 0–3 s; explicación clara; cierre calmado con solución."),
                ("Gráficos", "Badge dorado “Plazo legal”; countdown tipográfico; CTA pill verde."),
                ("CTA", "“Quiero hacer mi traspaso ahora” → WhatsApp con keyword TRASPASO."),
            ],
        },
        {
            "title": "Ejemplo 04 · Carrusel 7–8 slides — “Pasos que pide Medellín para visa de visitante”",
            "items": [
                ("Objetivo", "Autoridad local + guardados + shares a grupos de expats."),
                ("Tipo de plano", "Portada con foto ciudad + tipografía; interiores tipo guía editorial."),
                ("Ángulos", "Foto hero con profundidad (primer plano + fondo ciudad desenfocado)."),
                ("Visual", "Medellín recognoscible, iconos de requisitos, mock de expediente ordenado."),
                ("Ritmo", "Jerarquía tipográfica fija; 1 idea por slide; slide final CTA."),
                ("Gráficos", "Números grandes, verde + blanco, línea dorada sutil, logo discreto."),
                ("CTA", "“¿Quieres que revisemos tu caso? Escríbenos hoy”."),
            ],
        },
        {
            "title": "Ejemplo 05 · Reel 25–35 s — “Pasaporte EE.UU.: renovación sin enredos”",
            "items": [
                ("Objetivo", "Posicionar servicio de pasaportes con claridad y calma."),
                ("Tipo de plano", "Talking head + inserts DS-11/DS-82 (sin datos sensibles)."),
                ("Ángulos", "Frontal; detalle overhead de fotos 2×2 y formularios."),
                ("B-roll", "Pasaporte USA, fotos tipo documento, sobre/envío, check final."),
                ("Ritmo", "Explicativo elegante; cortes cada 3–4 s; sin saturación de texto."),
                ("Gráficos", "Etiquetas “Adultos” / “Menores”; checkmarks; end card bilingüe."),
                ("CTA", "“Cuéntanos tu caso — orientación sin compromiso”."),
            ],
        },
        {
            "title": "Ejemplo 06 · Reel 20–30 s — “Lo que nadie te dice del acompañamiento migratorio”",
            "items": [
                ("Objetivo", "Brand love + diferenciación emocional (valores de la firma)."),
                ("Tipo de plano", "Medio lateral ¾ + B-roll de atención al cliente / equipo."),
                ("Ángulos", "15–30°; travelling suave; luz más cálida."),
                ("B-roll", "WhatsApp en escritorio, reunión, sonrisa al cierre, sello/aprobación."),
                ("Ritmo", "Más aire; cuts emocionales; música suave con piano/ambient."),
                ("Gráficos", "Mínimos: subtítulos elegantes + frase final en serif."),
                ("CTA", "“Si quieres un proceso con alguien a tu lado, habla con nosotros”."),
            ],
        },
    ]

    for i, ex in enumerate(examples):
        if i == 3:
            page_break(doc)
            eyebrow(doc, "08 · Ejemplos de producción (cont.)")
            h1(doc, "Más referencias de ejecución")
        add_para(doc, ex["title"], size=11, bold=True, color=INK, space_before=10, space_after=4)
        for k, v in ex["items"]:
            add_rich(
                doc,
                [(f"{k}: ", {"bold": True, "color": PINK, "size": 9}), (v, {"size": 9, "color": BODY})],
                space_after=2,
            )

    page_break(doc)

    # ---- 09 PROCESO ----
    eyebrow(doc, "09 · Cómo trabajamos el mes")
    h1(doc, "Un flujo simple, pensado para producir")
    add_para(
        doc,
        "Sin fricción. Ustedes llegan con guiones y temas listos; yo me encargo de que cada pieza "
        "salga con nivel de marca.",
        size=11,
        color=MUTED,
    )
    steps = [
        ("01 · Kickoff & brief de rodaje", "Recibo los 12 guiones/temas del mes. Devuelvo un brief de producción por pieza: formato, planos, vestuario, B-roll y referencias de ritmo."),
        ("02 · Jornada(s) de grabación dirigida", "Sesión presencial o remota guiada. Optimizamos tomas, iluminación, audio y continuidad para editar en bloque."),
        ("03 · Edición creativa", "Corte, subtítulos, motion graphics, música, color y exportación 9:16. Priorizo hooks y retención sin perder el tono institucional."),
        ("04 · Revisión y entrega", "Una ronda de ajustes por pieza. Archivos finales listos para publicar (y portadas coherentes de Reel)."),
    ]
    for t, d in steps:
        add_para(doc, t, size=11, bold=True, color=PINK, space_before=8, space_after=2)
        add_para(doc, d, space_after=4)

    h2(doc, "Lo que necesito de ustedes")
    for t in [
        "Guiones y orden de prioridad del mes",
        "Acceso a logo / verdes de marca / tipografías si existen",
        "Disponibilidad para sesión de grabación",
        "Aprobador único de feedback (para velocidad)",
    ]:
        bullet(doc, t)

    h2(doc, "Lo que NO incluye este paquete")
    for t in [
        "Investigación o estrategia de contenidos",
        "Redacción de guiones desde cero",
        "Community management / respuestas",
        "Pauta paga o gestión de anuncios",
    ]:
        bullet(doc, t)

    page_break(doc)

    # ---- 10 CIERRE ----
    eyebrow(doc, "10 · Inversión y siguiente paso")
    h1(doc, "Listos para que cada publicación se sienta premium")
    add_para(doc, "PAQUETE MENSUAL · CREACIÓN DE CONTENIDO", size=8, bold=True, color=PINK, space_after=2)
    add_para(doc, "$600.000 COP / mes", size=26, color=INK, serif=True, space_after=4)
    add_para(
        doc,
        "12 publicaciones (Reels ≤ 40 s y/o carruseles) · Dirección de grabación · "
        "Producción audiovisual · Edición creativa · 1 ronda de ajustes por pieza",
        size=9,
        color=MUTED,
    )
    h2(doc, "Por qué esta inversión tiene sentido para VisasColTramit")
    add_para(
        doc,
        "En categorías de alto ticket emocional (visas, residencia, pasaportes), la gente no compra "
        "el precio más bajo: compra la sensación de estar en buenas manos. Un feed inconsistente "
        "genera duda. Un feed dirigido, iluminado y editado con rigor genera autoridad — y autoridad "
        "abre conversaciones de WhatsApp.",
    )
    h2(doc, "Siguiente paso")
    add_para(
        doc,
        "Si esta propuesta resuena, agendamos kickoff, definimos la primera jornada de grabación "
        "y arrancamos el mes 1 con sus guiones ya listos.",
    )
    add_para(doc, "Paola Hoyos Cardona", size=12, bold=True, color=INK, space_before=10, space_after=2)
    add_para(doc, "Comunicadora Social · Content Strategist · Producción de contenido", size=9, color=MUTED, space_after=2)
    add_para(doc, "Email: pahoyoscardona@gmail.com", size=10, space_after=1)
    add_para(doc, "Instagram: @paolaahoyosc", size=10, space_after=1)
    add_para(doc, "WhatsApp: disponible vía enlace de contacto", size=10, space_after=12)
    quote(
        doc,
        "Sus guiones ya tienen el mensaje. Yo me encargo de que se vean a la altura de la confianza que prometen.",
    )
    add_para(
        doc,
        "VisasColTramit × Paola Hoyos  ·  Julio 2026  ·  Documento comercial confidencial",
        size=8,
        color=MUTED,
        space_before=16,
    )

    out = "/workspace/propuestas/visascoltramit/Propuesta_Produccion_Contenido_VisasColTramit_PaolaHoyos.docx"
    doc.save(out)
    print("Saved", out)


if __name__ == "__main__":
    build()
